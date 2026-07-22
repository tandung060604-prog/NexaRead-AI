from __future__ import annotations

import json
import sys
import time
from collections.abc import Callable
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import TypedDict

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document as OpenXmlDocument  # noqa: E402
from ebooklib import epub  # type: ignore[import-untyped]  # noqa: E402

from app.services.docx_parser import parse_docx  # noqa: E402
from app.services.epub_parser import parse_epub  # noqa: E402
from app.services.normalized_document import NormalizedDocument  # noqa: E402
from app.services.web_parser import parse_web_page  # noqa: E402


class ParserCase(TypedDict):
    id: str
    passed: bool
    checks: dict[str, bool]
    latency_ms: float


def docx_fixture() -> bytes:
    document = OpenXmlDocument()
    document.core_properties.title = "Parser benchmark"
    document.add_heading("Chapter", 1)
    document.add_paragraph("Body paragraph")
    document.add_paragraph("List entry", style="List Bullet")
    table = document.add_table(1, 2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def epub_fixture() -> bytes:
    book = epub.EpubBook()
    book.set_identifier("parser-benchmark")
    book.set_title("Parser benchmark")
    book.set_language("en")
    chapter = epub.EpubHtml(title="Chapter", file_name="chapter.xhtml", lang="en")
    chapter.content = "<h1>Chapter</h1><p>Body paragraph</p><pre>code()</pre>"
    book.add_item(chapter)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]
    output = BytesIO()
    epub.write_epub(output, book)
    return output.getvalue()


def main() -> None:
    cases: list[ParserCase] = []
    started_all = time.perf_counter()
    fixtures: tuple[tuple[str, Callable[[], NormalizedDocument], list[str]], ...] = (
        (
            "docx",
            lambda: parse_docx(docx_fixture()),
            ["HEADING_1", "PARAGRAPH", "LIST_ITEM", "TABLE"],
        ),
        ("epub", lambda: parse_epub(epub_fixture()), ["HEADING_1", "PARAGRAPH", "CODE"]),
        (
            "url",
            lambda: parse_web_page(
                b"<html><head><title>Benchmark</title></head><body><main>"
                b"<h1>Chapter</h1><p>Body paragraph</p></main></body></html>",
                "https://example.com/article",
            ),
            ["HEADING_1", "PARAGRAPH"],
        ),
    )
    for name, parse, expected_types in fixtures:
        started = time.perf_counter()
        parsed = parse()
        actual_types = [block.block_type.value for block in parsed.blocks]
        checks = {
            "block_order": actual_types == expected_types,
            "metadata": bool(parsed.metadata),
            "content_non_empty": all(bool(block.text) for block in parsed.blocks),
            "stable_sequences": [block.sequence_number for block in parsed.blocks]
            == list(range(1, len(parsed.blocks) + 1)),
        }
        cases.append(
            {
                "id": name,
                "passed": all(checks.values()),
                "checks": checks,
                "latency_ms": round((time.perf_counter() - started) * 1000, 3),
            }
        )
    latencies = sorted(float(case["latency_ms"]) for case in cases)
    report = {
        "dataset_version": "2026.07-m7-v1",
        "generated_at": datetime.now(UTC).isoformat(),
        "case_count": len(cases),
        "passed": sum(bool(case["passed"]) for case in cases),
        "metrics": {
            "success_rate": round(sum(bool(case["passed"]) for case in cases) / len(cases), 4),
            "block_order_accuracy": round(
                sum(bool(dict(case["checks"])["block_order"]) for case in cases) / len(cases),
                4,
            ),
            "latency_p50_ms": latencies[len(latencies) // 2],
            "latency_max_ms": max(latencies),
            "total_latency_ms": round((time.perf_counter() - started_all) * 1000, 3),
        },
        "cases": cases,
    }
    output = ROOT / "evaluation" / "reports" / "parser-report.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report["metrics"], ensure_ascii=False))
    raise SystemExit(0 if report["passed"] == report["case_count"] else 1)


if __name__ == "__main__":
    main()
