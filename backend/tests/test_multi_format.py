from __future__ import annotations

from io import BytesIO
from uuid import UUID
from zipfile import ZIP_DEFLATED, ZipFile

import httpx
import pytest
from docx import Document as OpenXmlDocument
from ebooklib import epub  # type: ignore[import-untyped]

from app.services.docx_parser import parse_docx
from app.services.epub_parser import parse_epub
from app.services.normalized_document import BlockType, DocumentParseError
from app.services.processing import process_document_version
from app.services.url_import import FetchedUrl, UrlImportError, fetch_url, validate_remote_url
from app.services.web_parser import parse_web_page
from tests.conftest import ApiTestContext


def make_docx() -> bytes:
    document = OpenXmlDocument()
    document.core_properties.title = "DOCX Fixture"
    document.core_properties.author = "NexaRead Tests"
    document.add_heading("Architecture", level=1)
    document.add_paragraph("A paragraph that remains in source order.")
    document.add_paragraph("First list item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Latency"
    table.cell(1, 1).text = "20 ms"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_epub() -> bytes:
    book = epub.EpubBook()
    book.set_identifier("nexaread-test")
    book.set_title("EPUB Fixture")
    book.set_language("en")
    book.add_author("NexaRead Tests")
    first = epub.EpubHtml(title="First", file_name="first.xhtml", lang="en")
    first.content = "<h1>First chapter</h1><p>Readable first chapter text.</p>"
    second = epub.EpubHtml(title="Second", file_name="second.xhtml", lang="en")
    second.content = "<h2>Second section</h2><pre>print('safe')</pre>"
    book.add_item(first)
    book.add_item(second)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.toc = (first, second)
    book.spine = ["nav", first, second]
    output = BytesIO()
    epub.write_epub(output, book)
    return output.getvalue()


def test_docx_parser_preserves_structure_and_order() -> None:
    parsed = parse_docx(make_docx())

    assert parsed.metadata["title"] == "DOCX Fixture"
    assert [block.block_type for block in parsed.blocks] == [
        BlockType.HEADING_1,
        BlockType.PARAGRAPH,
        BlockType.LIST_ITEM,
        BlockType.TABLE,
    ]
    assert parsed.blocks[-1].metadata["rows"] == [
        ["Name", "Value"],
        ["Latency", "20 ms"],
    ]


def test_epub_parser_uses_spine_order_and_metadata() -> None:
    parsed = parse_epub(make_epub())

    assert parsed.metadata["title"] == "EPUB Fixture"
    assert parsed.page_count == 2
    assert [page.page_number for page in parsed.pages] == [1, 2]
    assert parsed.blocks[0].text == "First chapter"
    assert parsed.blocks[-1].block_type == BlockType.CODE


def test_epub_parser_refuses_drm() -> None:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "META-INF/encryption.xml",
            '<encryption><EncryptedData><EncryptionMethod Algorithm="urn:drm"/>'
            "</EncryptedData></encryption>",
        )
    with pytest.raises(DocumentParseError) as error:
        parse_epub(output.getvalue())
    assert error.value.code == "EPUB_DRM_PROTECTED"


def test_web_parser_extracts_main_content_and_sanitizes_active_content() -> None:
    html = b"""
    <html><head><title>Safe article</title><script>alert('x')</script></head>
    <body><nav>Navigation noise</nav><article><h1>Heading</h1>
    <p>Useful article paragraph.</p><img src="/cover.png" onerror="alert(1)"></article></body>
    </html>
    """
    parsed = parse_web_page(html, "https://example.com/article")

    assert parsed.metadata["title"] == "Safe article"
    assert "Navigation noise" not in parsed.pages[0].text
    assert "alert" not in parsed.pages[0].text
    assert parsed.assets[0].source_reference == "https://example.com/cover.png"


@pytest.mark.anyio
@pytest.mark.parametrize(
    "url",
    [
        "file:///etc/passwd",
        "http://localhost/admin",
        "http://127.0.0.1/",
        "http://169.254.169.254/latest/meta-data/",
        "http://10.0.0.1/",
        "http://example.com:8080/",
    ],
)
async def test_url_validation_rejects_unsafe_targets(url: str) -> None:
    with pytest.raises(UrlImportError):
        await validate_remote_url(url)


@pytest.mark.anyio
async def test_url_fetch_revalidates_redirect_target() -> None:
    async def public_resolver(_: str, __: int) -> list[str]:
        return ["93.184.216.34"]

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(302, headers={"location": "http://127.0.0.1/private"})

    with pytest.raises(UrlImportError, match="private or unsafe"):
        await fetch_url(
            "https://example.com/start",
            max_bytes=1024,
            timeout_seconds=2,
            max_redirects=2,
            resolver=public_resolver,
            transport=httpx.MockTransport(handler),
        )


@pytest.mark.anyio
async def test_uploads_and_processes_docx(api_context: ApiTestContext) -> None:
    response = await api_context.client.post(
        "/api/documents/upload",
        files={
            "file": (
                "guide.docx",
                make_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    payload = response.json()
    assert payload["source_type"] == "docx"
    assert payload["mime_type"].endswith("wordprocessingml.document")
    assert api_context.session_factory is not None

    await process_document_version(
        UUID(payload["id"]),
        UUID(payload["versions"][0]["id"]),
        UUID(payload["processing_jobs"][0]["id"]),
        storage=api_context.storage,
        session_factory=api_context.session_factory,
    )
    detail = await api_context.client.get(f"/api/documents/{payload['id']}")
    blocks = await api_context.client.get(f"/api/documents/{payload['id']}/blocks?limit=100")
    assert detail.json()["status"] == "AI_READY"
    assert detail.json()["layout_type"] == "GENERAL_DOCUMENT"
    assert blocks.json()["total"] == 4


@pytest.mark.anyio
async def test_import_url_api_queues_sanitized_source(
    api_context: ApiTestContext, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def fake_fetch(*_: object, **__: object) -> FetchedUrl:
        return FetchedUrl(
            b"<html><head><title>Imported</title></head><body><main><h1>Web</h1>"
            b"<p>Grounded web content.</p></main></body></html>",
            "https://example.com/article",
            "text/html",
        )

    monkeypatch.setattr("app.services.url_documents.fetch_url", fake_fetch)
    collection = await api_context.client.post(
        "/api/collections",
        json={"name": "Web research"},
    )
    response = await api_context.client.post(
        "/api/documents/import-url",
        json={
            "url": "https://example.com/article",
            "document_type_override": "technical",
            "collection_id": collection.json()["id"],
        },
    )
    assert response.status_code == 201
    payload = response.json()
    assert payload["source_type"] == "url"
    assert payload["source_url"] == "https://example.com/article"
    assert payload["document_type_override"] == "TECHNICAL"
    assert payload["layout_type"] == "TECHNICAL_DOCUMENTATION"
    assert payload["collection_id"] == collection.json()["id"]
    assert len(api_context.queue.messages) == 1
